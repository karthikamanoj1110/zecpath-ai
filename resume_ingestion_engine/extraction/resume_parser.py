import re
import os
import sys

from datetime import datetime
from resume_ingestion_engine.extraction.skill_extractor import extract_resume_skills

def convert_to_txt(input_path: str) -> str:
    """
    Convert a PDF or DOCX resume to a .txt file.
    Returns the path to the saved .txt file.
    """
    ext = os.path.splitext(input_path)[1].lower()
    os.makedirs('extraction_output', exist_ok=True)

    # Output filename: same name, .txt extension
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    txt_path  = f'extraction_output/{base_name}.txt'

    if ext == '.pdf':
        try:
            import pdfplumber
            with pdfplumber.open(input_path) as pdf:
                text = "\n".join(
                    page.extract_text() or "" for page in pdf.pages
                )
        except ImportError:
            raise ImportError("Install pdfplumber:  pip install pdfplumber")

    elif ext in ('.docx', '.doc'):
        try:
            import docx
            doc  = docx.Document(input_path)
            text = "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            raise ImportError("Install python-docx:  pip install python-docx")

    elif ext == '.txt':
        # Already a txt — just copy it over
        with open(input_path, 'r', encoding='utf-8') as f:
            text = f.read()

    else:
        raise ValueError(f"Unsupported file type: {ext}")

    # Save as .txt
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(text)

    print(f"📄 Converted '{os.path.basename(input_path)}' → '{txt_path}'")
    return txt_path

def parse_resume(text:str)->dict:
    if not text:
        return{"skills":[]}
    
    text = re.sub(r"\s+"," ",text).strip()
    skills = extract_resume_skills(text)

    return{
        "skills": skills
    }


def extract_text_from_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def clean_text(text):
    text = ' '.join(text.split())
    text = re.sub(r'[•·●]', '•', text)
    return text

def mask_privacy(text):
    text = re.sub(r'\b[\w\.-]+@[\w\.-]+\.\w+\b', '[EMAIL]', text)
    text =  re.sub(r'(\+?\d{1,3}[\s\-.]?)?\(?\d{2,5}\)?[\s\-.]?\d{3,5}[\s\-.]?\d{4,5}', '[PHONE]', text)
    return text

def save_output(text, filename):
    os.makedirs('extraction_output', exist_ok=True)
    path = f'extraction_output/{filename}'
    with open(path, 'w', encoding='utf-8') as f:
        f.write(text)
    return path

def main():
    print("="*70)
    print("RESUME TEXT EXTRACTION ENGINE - DAY 5")
    print("="*70)
    
    # ✅ NEW: Accept resume path from command line, convert to .txt
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
        test_file  = convert_to_txt(input_file)   # ← auto-converts & saves .txt
    else:
        test_file  = 'my_resume.txt'               # ← fallback default
        print(f"ℹ️  No file given, using default: {test_file}")
    
    # Step 1: Extract
    raw_text = extract_text_from_file(test_file)
    print(f"\n📄 Original: {len(raw_text)} chars")
    
    # Step 2: Clean
    cleaned = clean_text(raw_text)
    
    # Step 3: Mask privacy
    masked = mask_privacy(cleaned)
    print(f"✅ Cleaned : {len(masked)} chars")

     # Step 4: Parse skills  ✅ NEW — actually calls parse_resume now
    result = parse_resume(masked)
    print(f"🧠 Skills  : {result['skills']}")
    
    # Step 5: Save
    base_name   = os.path.splitext(os.path.basename(test_file))[0]
    output_path = save_output(masked, f'cleaned_{base_name}.txt')
    print(f"💾 Saved to: {output_path}")
    
    # Step 6: Report
    report = f"""
📊 PROCESSING REPORT
-------------------
File: {test_file}
Original chars: {len(raw_text)}
Cleaned chars : {len(masked)}
Reduction     : {len(raw_text)-len(masked)} chars
Time          : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""
    
    with open('extraction_output/EXTRACTION_REPORT.txt', 'w',encoding='utf-8') as f:
        f.write(report)
    print("📄 Report saved")
    
    print("\n" + "="*70)
    print("✅ DAY 5 COMPLETED")
    print("="*70)

import os

def read_input_file(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    
    elif ext in (".docx", ".doc"):
        import docx
        doc = docx.Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)
    
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    
    else:
        raise ValueError(f"Unsupported file type: {ext}")


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python -m resume_ingestion_engine.extraction.resume_parser <resume_file>")
        sys.exit(1)

    text   = read_input_file(sys.argv[1])
    result = parse_resume(text)
    print(result)