from pathlib import Path

import pdfplumber
from docx import Document


def read_resume(file_path):
    """
    Entry point.
    Automatically detects PDF or DOCX and returns layout-aware blocks.
    """

    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return _read_pdf(file_path)

    elif ext == ".docx":
        return _read_docx(file_path)

    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ---------------- PDF ----------------

def _read_pdf(file_path):
    blocks = []

    with pdfplumber.open(file_path) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):

            words = page.extract_words(
                use_text_flow=True,
                keep_blank_chars=False,
                extra_attrs=["size"]
            )

            for w in words:
                blocks.append({
                    "type": "text",
                    "source": "pdf",
                    "page": page_index,
                    "text": w["text"],
                    "bbox": [w["x0"], w["top"], w["x1"], w["bottom"]],
                    "font_size": w.get("size")
                })

    return blocks


# ---------------- DOCX ----------------

def _read_docx(file_path):
    doc = Document(file_path)

    blocks = []

    # paragraphs
    for p in doc.paragraphs:

        text = p.text.strip()
        if not text:
            continue

        blocks.append({
            "type": "paragraph",
            "source": "docx",
            "text": text,
            "style": p.style.name if p.style else None,
            "is_bold": any(run.bold for run in p.runs),
            "is_italic": any(run.italic for run in p.runs)
        })

    # tables
    for table in doc.tables:

        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])

        blocks.append({
            "type": "table",
            "source": "docx",
            "rows": rows
        })

    return blocks
